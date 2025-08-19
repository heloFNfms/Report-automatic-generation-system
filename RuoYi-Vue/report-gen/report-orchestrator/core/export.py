import os
import io
import tempfile
from typing import Dict, Any, Optional
from datetime import datetime
import asyncio
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import weasyprint
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    import boto3
    from botocore.exceptions import ClientError
    BOTO3_AVAILABLE = True
except ImportError:
    BOTO3_AVAILABLE = False

try:
    from oss2 import Auth, Bucket
    OSS2_AVAILABLE = True
except ImportError:
    OSS2_AVAILABLE = False

from .logger import logger

class ExportManager:
    """导出管理器"""
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "report_exports"
        self.temp_dir.mkdir(exist_ok=True)
        
        # OSS 配置
        self.oss_config = {
            'aliyun': {
                'access_key_id': os.getenv('ALIYUN_OSS_ACCESS_KEY_ID'),
                'access_key_secret': os.getenv('ALIYUN_OSS_ACCESS_KEY_SECRET'),
                'endpoint': os.getenv('ALIYUN_OSS_ENDPOINT'),
                'bucket_name': os.getenv('ALIYUN_OSS_BUCKET_NAME')
            },
            'minio': {
                'endpoint': os.getenv('MINIO_ENDPOINT', 'http://localhost:9000'),
                'access_key': os.getenv('MINIO_ACCESS_KEY'),
                'secret_key': os.getenv('MINIO_SECRET_KEY'),
                'bucket_name': os.getenv('MINIO_BUCKET_NAME', 'reports')
            }
        }
    
    async def export_to_word(self, task_data: Dict[str, Any]) -> Dict[str, Any]:
        """导出为 Word 文档"""
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not available. Please install python-docx."}
        
        try:
            logger.info(f"Starting Word export for task {task_data.get('task_id')}")
            
            # 创建文档
            doc = Document()
            
            # 设置文档样式
            self._setup_word_styles(doc)
            
            # 添加标题
            title = task_data.get('project_name', 'Research Report')
            doc.add_heading(title, 0)
            
            # 添加元信息
            self._add_word_metadata(doc, task_data)
            
            # 添加大纲
            if 'outline' in task_data:
                self._add_word_outline(doc, task_data['outline'])
            
            # 添加内容
            if 'sections' in task_data:
                self._add_word_content(doc, task_data['sections'])
            
            # 添加最终报告
            if 'final_report' in task_data:
                self._add_word_final_report(doc, task_data['final_report'])
            
            # 保存文档
            filename = f"report_{task_data.get('task_id', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = self.temp_dir / filename
            
            doc.save(str(file_path))
            
            logger.info(f"Word export completed: {filename}")
            
            return {
                "status": "success",
                "format": "docx",
                "filename": filename,
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size
            }
        except Exception as e:
            logger.error(f"Word export failed: {str(e)}")
            return {"error": f"Word export failed: {str(e)}"}
    
    async def export_to_pdf(self, task_data: Dict[str, Any]) -> Dict[str, Any]:
        """导出为 PDF 文档"""
        if not WEASYPRINT_AVAILABLE:
            return {"error": "WeasyPrint not available. Please install WeasyPrint."}
        
        try:
            logger.info(f"Starting PDF export for task {task_data.get('task_id')}")
            
            # 生成 HTML 内容
            html_content = self._generate_html_content(task_data)
            
            # 生成 CSS 样式
            css_content = self._generate_css_styles()
            
            # 创建 PDF
            filename = f"report_{task_data.get('task_id', 'unknown')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = self.temp_dir / filename
            
            html_doc = HTML(string=html_content)
            css_doc = CSS(string=css_content)
            
            html_doc.write_pdf(str(file_path), stylesheets=[css_doc])
            
            logger.info(f"PDF export completed: {filename}")
            
            return {
                "status": "success",
                "format": "pdf",
                "filename": filename,
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size
            }
        except Exception as e:
            logger.error(f"PDF export failed: {str(e)}")
            return {"error": f"PDF export failed: {str(e)}"}
    
    async def upload_to_oss(self, file_path: str, provider: str = 'minio') -> Dict[str, Any]:
        """上传文件到对象存储"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {"error": "File not found"}
            
            if provider == 'aliyun':
                return await self._upload_to_aliyun_oss(file_path)
            elif provider == 'minio':
                return await self._upload_to_minio(file_path)
            else:
                return {"error": f"Unsupported provider: {provider}"}
        except Exception as e:
            logger.error(f"OSS upload failed: {str(e)}")
            return {"error": f"OSS upload failed: {str(e)}"}
    
    async def _upload_to_aliyun_oss(self, file_path: Path) -> Dict[str, Any]:
        """上传到阿里云 OSS"""
        if not OSS2_AVAILABLE:
            return {"error": "oss2 not available. Please install oss2."}
        
        config = self.oss_config['aliyun']
        if not all([config['access_key_id'], config['access_key_secret'], config['endpoint'], config['bucket_name']]):
            return {"error": "Aliyun OSS configuration incomplete"}
        
        try:
            auth = Auth(config['access_key_id'], config['access_key_secret'])
            bucket = Bucket(auth, config['endpoint'], config['bucket_name'])
            
            # 生成对象键
            object_key = f"reports/{datetime.now().strftime('%Y/%m/%d')}/{file_path.name}"
            
            # 上传文件
            with open(file_path, 'rb') as f:
                result = bucket.put_object(object_key, f)
            
            # 生成访问 URL
            url = f"https://{config['bucket_name']}.{config['endpoint'].replace('https://', '').replace('http://', '')}/{object_key}"
            
            logger.info(f"File uploaded to Aliyun OSS: {object_key}")
            
            return {
                "status": "success",
                "provider": "aliyun",
                "object_key": object_key,
                "url": url,
                "etag": result.etag
            }
        except Exception as e:
            return {"error": f"Aliyun OSS upload failed: {str(e)}"}
    
    async def _upload_to_minio(self, file_path: Path) -> Dict[str, Any]:
        """上传到 MinIO"""
        if not BOTO3_AVAILABLE:
            return {"error": "boto3 not available. Please install boto3."}
        
        config = self.oss_config['minio']
        if not all([config['endpoint'], config['access_key'], config['secret_key']]):
            return {"error": "MinIO configuration incomplete"}
        
        try:
            # 创建 S3 客户端（MinIO 兼容 S3 API）
            s3_client = boto3.client(
                's3',
                endpoint_url=config['endpoint'],
                aws_access_key_id=config['access_key'],
                aws_secret_access_key=config['secret_key']
            )
            
            # 确保存储桶存在
            try:
                s3_client.head_bucket(Bucket=config['bucket_name'])
            except ClientError:
                s3_client.create_bucket(Bucket=config['bucket_name'])
            
            # 生成对象键
            object_key = f"reports/{datetime.now().strftime('%Y/%m/%d')}/{file_path.name}"
            
            # 上传文件
            with open(file_path, 'rb') as f:
                s3_client.upload_fileobj(f, config['bucket_name'], object_key)
            
            # 生成访问 URL
            url = f"{config['endpoint']}/{config['bucket_name']}/{object_key}"
            
            logger.info(f"File uploaded to MinIO: {object_key}")
            
            return {
                "status": "success",
                "provider": "minio",
                "object_key": object_key,
                "url": url
            }
        except Exception as e:
            return {"error": f"MinIO upload failed: {str(e)}"}
    
    def _setup_word_styles(self, doc: Document):
        """设置 Word 文档样式"""
        # 设置正文样式
        styles = doc.styles
        
        # 标题样式
        if 'Custom Heading 1' not in [s.name for s in styles]:
            heading1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = 'Arial'
            heading1_style.font.size = Pt(16)
            heading1_style.font.bold = True
        
        # 正文样式
        if 'Custom Normal' not in [s.name for s in styles]:
            normal_style = styles.add_style('Custom Normal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = 'Arial'
            normal_style.font.size = Pt(11)
    
    def _add_word_metadata(self, doc: Document, task_data: Dict[str, Any]):
        """添加文档元信息"""
        # 添加元信息表格
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # 填充元信息
        cells = table.rows[0].cells
        cells[0].text = '项目名称'
        cells[1].text = task_data.get('project_name', 'N/A')
        
        cells = table.rows[1].cells
        cells[0].text = '公司名称'
        cells[1].text = task_data.get('company_name', 'N/A')
        
        cells = table.rows[2].cells
        cells[0].text = '生成时间'
        cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        cells = table.rows[3].cells
        cells[0].text = '任务ID'
        cells[1].text = task_data.get('task_id', 'N/A')
        
        doc.add_paragraph()  # 添加空行
    
    def _add_word_outline(self, doc: Document, outline_data: Dict[str, Any]):
        """添加大纲到 Word 文档"""
        doc.add_heading('研究大纲', level=1)
        
        if isinstance(outline_data, dict) and 'outline' in outline_data:
            outline_text = outline_data['outline']
        else:
            outline_text = str(outline_data)
        
        # 分段添加大纲内容
        paragraphs = outline_text.split('\n')
        for para in paragraphs:
            if para.strip():
                if para.strip().startswith('#'):
                    # 处理标题
                    level = para.count('#')
                    title = para.replace('#', '').strip()
                    doc.add_heading(title, level=min(level, 3))
                else:
                    doc.add_paragraph(para.strip())
        
        doc.add_page_break()
    
    def _add_word_content(self, doc: Document, sections_data: list):
        """添加章节内容到 Word 文档"""
        doc.add_heading('详细内容', level=1)
        
        for i, section in enumerate(sections_data, 1):
            if isinstance(section, dict):
                title = section.get('title', f'第{i}章')
                content = section.get('content', '')
            else:
                title = f'第{i}章'
                content = str(section)
            
            doc.add_heading(title, level=2)
            
            # 分段添加内容
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            doc.add_paragraph()  # 添加空行
    
    def _add_word_final_report(self, doc: Document, final_report: str):
        """添加最终报告到 Word 文档"""
        doc.add_page_break()
        doc.add_heading('最终报告', level=1)
        
        # 分段添加最终报告内容
        paragraphs = final_report.split('\n')
        for para in paragraphs:
            if para.strip():
                if para.strip().startswith('#'):
                    # 处理标题
                    level = para.count('#')
                    title = para.replace('#', '').strip()
                    doc.add_heading(title, level=min(level + 1, 3))
                else:
                    doc.add_paragraph(para.strip())
    
    def _generate_html_content(self, task_data: Dict[str, Any]) -> str:
        """生成 HTML 内容"""
        html = f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{task_data.get('project_name', 'Research Report')}</title>
</head>
<body>
    <div class="container">
        <header>
            <h1 class="main-title">{task_data.get('project_name', 'Research Report')}</h1>
            <div class="metadata">
                <p><strong>公司名称：</strong>{task_data.get('company_name', 'N/A')}</p>
                <p><strong>生成时间：</strong>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
                <p><strong>任务ID：</strong>{task_data.get('task_id', 'N/A')}</p>
            </div>
        </header>
        
        <main>
"""
        
        # 添加大纲
        if 'outline' in task_data:
            html += "<section class='outline'>\n<h2>研究大纲</h2>\n"
            outline_data = task_data['outline']
            if isinstance(outline_data, dict) and 'outline' in outline_data:
                outline_text = outline_data['outline']
            else:
                outline_text = str(outline_data)
            
            html += f"<div class='outline-content'>{self._format_text_to_html(outline_text)}</div>\n"
            html += "</section>\n"
        
        # 添加章节内容
        if 'sections' in task_data:
            html += "<section class='content'>\n<h2>详细内容</h2>\n"
            for i, section in enumerate(task_data['sections'], 1):
                if isinstance(section, dict):
                    title = section.get('title', f'第{i}章')
                    content = section.get('content', '')
                else:
                    title = f'第{i}章'
                    content = str(section)
                
                html += f"<div class='section'>\n<h3>{title}</h3>\n"
                html += f"<div class='section-content'>{self._format_text_to_html(content)}</div>\n"
                html += "</div>\n"
            html += "</section>\n"
        
        # 添加最终报告
        if 'final_report' in task_data:
            html += "<section class='final-report'>\n<h2>最终报告</h2>\n"
            html += f"<div class='final-content'>{self._format_text_to_html(task_data['final_report'])}</div>\n"
            html += "</section>\n"
        
        html += """
        </main>
    </div>
</body>
</html>
"""
        
        return html
    
    def _format_text_to_html(self, text: str) -> str:
        """将文本格式化为 HTML"""
        if not text:
            return ""
        
        # 转义 HTML 特殊字符
        import html
        text = html.escape(text)
        
        # 处理换行
        paragraphs = text.split('\n')
        formatted_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            if para:
                if para.startswith('#'):
                    # 处理标题
                    level = para.count('#')
                    title = para.replace('#', '').strip()
                    formatted_paragraphs.append(f'<h{min(level + 2, 6)}>{title}</h{min(level + 2, 6)}>')
                else:
                    formatted_paragraphs.append(f'<p>{para}</p>')
        
        return '\n'.join(formatted_paragraphs)
    
    def _generate_css_styles(self) -> str:
        """生成 CSS 样式"""
        return """
@page {
    size: A4;
    margin: 2cm;
}

body {
    font-family: 'SimSun', 'Arial', sans-serif;
    font-size: 12pt;
    line-height: 1.6;
    color: #333;
}

.container {
    max-width: 100%;
}

header {
    text-align: center;
    margin-bottom: 2cm;
    border-bottom: 2px solid #333;
    padding-bottom: 1cm;
}

.main-title {
    font-size: 24pt;
    font-weight: bold;
    margin-bottom: 1cm;
    color: #2c3e50;
}

.metadata {
    font-size: 11pt;
    color: #666;
}

.metadata p {
    margin: 0.2cm 0;
}

h2 {
    font-size: 16pt;
    font-weight: bold;
    color: #2c3e50;
    margin-top: 1.5cm;
    margin-bottom: 0.5cm;
    border-bottom: 1px solid #bdc3c7;
    padding-bottom: 0.2cm;
}

h3 {
    font-size: 14pt;
    font-weight: bold;
    color: #34495e;
    margin-top: 1cm;
    margin-bottom: 0.3cm;
}

h4, h5, h6 {
    font-size: 12pt;
    font-weight: bold;
    color: #34495e;
    margin-top: 0.8cm;
    margin-bottom: 0.2cm;
}

p {
    margin: 0.3cm 0;
    text-align: justify;
    text-indent: 2em;
}

.section {
    margin-bottom: 1.5cm;
}

.outline-content, .section-content, .final-content {
    margin-left: 1cm;
}

.final-report {
    page-break-before: always;
}
"""
    
    async def cleanup_temp_files(self, max_age_hours: int = 24):
        """清理临时文件"""
        try:
            current_time = datetime.now()
            for file_path in self.temp_dir.glob('*'):
                if file_path.is_file():
                    file_age = current_time - datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_age.total_seconds() > max_age_hours * 3600:
                        file_path.unlink()
                        logger.info(f"Cleaned up temp file: {file_path.name}")
        except Exception as e:
            logger.error(f"Temp file cleanup failed: {str(e)}")

# 全局导出管理器实例
export_manager = ExportManager()